from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parent
TENANT_A = ROOT / "tenant-a"
TENANT_B = ROOT / "tenant-b"
HEADERS = [
    "Date",
    "Employee ID",
    "Employee Name",
    "Status",
    "Department",
    "Check In",
    "Check Out",
    "Total Hours",
    "Attendance %",
]
ROWS = [
    [
        "2026-08-01",
        "SYN-001",
        "Asha Fiction",
        "Present",
        "Engineering",
        "09:00",
        "17:00",
        "8",
        "100",
    ],
    ["2026-08-01", "SYN-002", "Mira Example", "Absent", "Support", "", "", "0", "0"],
    ["2026-08-02", "SYN-003", "Noor Sample", "Leave", "Finance", "", "", "", "100"],
    ["2026-08-02", "SYN-004", "Ira Demo", "WFH", "Engineering", "09:15", "17:15", "8", "100"],
]


def csv_text(rows: list[list[str]]) -> str:
    return "\n".join(",".join(row) for row in [HEADERS, *rows]) + "\n"


def create_csv() -> None:
    (TENANT_A / "attendance.csv").write_text(csv_text(ROWS), encoding="utf-8")
    tenant_b_rows = [
        [
            "2026-08-01",
            "SYN-B01",
            "Ravi Imaginary",
            "Present",
            "Operations",
            "08:30",
            "16:30",
            "8",
            "100",
        ],
        [
            "2026-08-02",
            "SYN-B02",
            "Tara Fiction",
            "Late",
            "Operations",
            "09:20",
            "17:00",
            "7.67",
            "96",
        ],
    ]
    (TENANT_B / "attendance.csv").write_text(csv_text(tenant_b_rows), encoding="utf-8")


def create_xlsx() -> None:
    workbook = Workbook()
    engineering = workbook.active
    engineering.title = "Engineering"
    engineering.append(HEADERS)
    for row in (ROWS[0], ROWS[3]):
        engineering.append(row)
    other = workbook.create_sheet("Other Departments")
    other.append(HEADERS)
    for row in (ROWS[1], ROWS[2]):
        other.append(row)
    workbook.save(TENANT_A / "attendance.xlsx")


def create_docx() -> None:
    document = Document()
    document.add_heading("Synthetic Attendance Evidence", level=1)
    document.add_paragraph("All names and employee identifiers in this document are fictional.")
    table = document.add_table(rows=1, cols=len(HEADERS))
    for index, header in enumerate(HEADERS):
        table.rows[0].cells[index].text = header
    for row in ROWS:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.save(TENANT_A / "attendance.docx")


def create_text_pdf() -> None:
    pdf = canvas.Canvas(str(TENANT_A / "attendance-text.pdf"))
    y = 800
    for line in csv_text(ROWS).splitlines():
        pdf.drawString(30, y, line)
        y -= 20
    pdf.save()


def create_scanned_inputs() -> None:
    width, height = 2100, 500
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"
    font = ImageFont.truetype(font_path, 28)
    lines = csv_text(ROWS[:2]).splitlines()
    for index, line in enumerate(lines):
        draw.text((25, 30 + index * 70), line, fill="black", font=font)
    # A slight blur makes OCR confidence realistically imperfect while leaving
    # the synthetic rows readable and traceable.
    image = image.filter(ImageFilter.GaussianBlur(radius=0.55))
    image_path = TENANT_A / "attendance-scan.png"
    image.save(image_path)

    output = BytesIO()
    image.save(output, format="PNG")
    pdf = canvas.Canvas(str(TENANT_A / "attendance-scanned.pdf"), pagesize=(900, 300))
    pdf.drawImage(ImageReader(BytesIO(output.getvalue())), 0, 0, width=900, height=214)
    pdf.save()


def main() -> None:
    TENANT_A.mkdir(parents=True, exist_ok=True)
    TENANT_B.mkdir(parents=True, exist_ok=True)
    create_csv()
    create_xlsx()
    create_docx()
    create_text_pdf()
    create_scanned_inputs()


if __name__ == "__main__":
    main()
