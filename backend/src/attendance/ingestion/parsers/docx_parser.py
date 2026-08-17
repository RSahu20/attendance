from io import BytesIO

from docx import Document as OpenDocument

from attendance.domain.attendance import ExtractionMethod
from attendance.ingestion.parsers.base import DocumentParser
from attendance.ingestion.parsers.tabular import table_units
from attendance.ingestion.types import IngestionError, ParsedUnit


class DOCXParser(DocumentParser):
    name = "docx"

    def parse(self, content: bytes, filename: str) -> list[ParsedUnit]:
        try:
            document = OpenDocument(BytesIO(content))
            units: list[ParsedUnit] = []
            for index, paragraph in enumerate(document.paragraphs, start=1):
                if paragraph.text.strip():
                    units.append(
                        ParsedUnit(
                            source_unit_key=f"paragraph:{index}",
                            unit_type="paragraph",
                            source_locator={"file": filename, "paragraph": index},
                            raw_text=paragraph.text.strip(),
                            extraction_method=ExtractionMethod.DOCX,
                        )
                    )
            for table_index, table in enumerate(document.tables, start=1):
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                try:
                    units.extend(
                        table_units(
                            rows,
                            filename=filename,
                            location={"table": table_index},
                            key_prefix=f"table:{table_index}",
                            method=ExtractionMethod.DOCX,
                        )
                    )
                except ValueError:
                    for row_index, row in enumerate(rows, start=1):
                        units.append(
                            ParsedUnit(
                                source_unit_key=f"table:{table_index}:text-row:{row_index}",
                                unit_type="table_row",
                                source_locator={
                                    "file": filename,
                                    "table": table_index,
                                    "row": row_index,
                                },
                                raw_text=" | ".join(row),
                                extraction_method=ExtractionMethod.DOCX,
                            )
                        )
            if not units:
                raise ValueError("Document contains no readable content")
            return units
        except (ValueError, KeyError, OSError) as exc:
            raise IngestionError("unreadable_docx", "The DOCX file could not be parsed") from exc
