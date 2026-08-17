from datetime import date
from io import BytesIO

from docx import Document
from openpyxl import Workbook
from PIL import Image
from reportlab.pdfgen import canvas

from attendance.domain.attendance import ExtractionMethod, ReviewStatus
from attendance.ingestion.parsers.csv_parser import CSVParser
from attendance.ingestion.parsers.docx_parser import DOCXParser
from attendance.ingestion.parsers.ocr_parser import OCRParser
from attendance.ingestion.parsers.pdf_parser import PDFParser
from attendance.ingestion.parsers.xlsx_parser import XLSXParser
from attendance.ingestion.types import IngestionError
from attendance.providers.ocr.base import OCRProvider, OCRResult

HEADERS = ["Date", "Employee ID", "Employee Name", "Status", "Department"]
ROW = [date(2026, 8, 1), "EMP-001", "Asha Fiction", "Present", "Engineering"]


class StubOCRProvider(OCRProvider):
    def __init__(self, confidence: float = 0.91) -> None:
        self.confidence = confidence

    def extract(self, image: Image.Image) -> OCRResult:
        return OCRResult(
            text="Date,Employee ID,Employee Name,Status,Department\n"
            "2026-08-01,EMP-005,Noor Example,WFH,Support",
            confidence=self.confidence,
        )


def test_csv_parser_maps_aliases_and_preserves_row_number() -> None:
    content = (
        b"Attendance Date,emp_id,name,attendance,dept,check-in,check-out,hours,attendance %\n"
        b"2026-08-01,EMP-001,Asha Fiction,Present,Engineering,09:00,17:00,8,100\n"
    )

    units = CSVParser().parse(content, "sample.csv")

    assert units[0].structured_data["employee_id"] == "EMP-001"
    assert units[0].source_locator == {"file": "sample.csv", "row": 2}
    assert units[0].extraction_method == ExtractionMethod.CSV


def test_csv_parser_reports_missing_attendance_header() -> None:
    content = b"agent_id,agent_name,capability\nAG-001,Planner,orchestration\n"

    try:
        CSVParser().parse(content, "agent_registry.csv")
    except IngestionError as exc:
        assert exc.code == "no_attendance_header"
        assert "attendance header" in exc.safe_message
    else:
        raise AssertionError("A non-attendance CSV must not be accepted")


def test_xlsx_parser_supports_multiple_sheets_and_lineage() -> None:
    workbook = Workbook()
    first = workbook.active
    first.title = "Engineering"
    first.append(HEADERS)
    first.append(ROW)
    second = workbook.create_sheet("Support")
    second.append(HEADERS)
    second.append([date(2026, 8, 2), "EMP-002", "Mira Example", "Absent", "Support"])
    output = BytesIO()
    workbook.save(output)

    units = XLSXParser().parse(output.getvalue(), "sample.xlsx")

    assert len(units) == 2
    assert units[1].source_locator == {"file": "sample.xlsx", "sheet": "Support", "row": 2}


def test_docx_parser_extracts_paragraphs_tables_and_rows() -> None:
    document = Document()
    document.add_paragraph("Synthetic attendance evidence")
    table = document.add_table(rows=2, cols=len(HEADERS))
    for index, value in enumerate(HEADERS):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(ROW):
        table.rows[1].cells[index].text = str(value)
    output = BytesIO()
    document.save(output)

    units = DOCXParser().parse(output.getvalue(), "sample.docx")
    attendance = [unit for unit in units if unit.unit_type == "attendance_row"]

    assert any(unit.unit_type == "paragraph" for unit in units)
    assert attendance[0].source_locator == {"file": "sample.docx", "table": 1, "row": 2}


def test_text_pdf_parser_preserves_page_and_extracts_rows() -> None:
    output = BytesIO()
    pdf = canvas.Canvas(output)
    pdf.drawString(40, 800, "Date,Employee ID,Employee Name,Status,Department")
    pdf.drawString(40, 780, "2026-08-01,EMP-003,Ira Example,Leave,Finance")
    pdf.save()
    parser = PDFParser(OCRParser(StubOCRProvider(), confidence_threshold=0.8))

    units = parser.parse(output.getvalue(), "sample.pdf")

    assert units[0].source_locator["page"] == 1
    assert units[0].structured_data["employee_id"] == "EMP-003"
    assert units[0].extraction_method == ExtractionMethod.PDF_TEXT


def test_ocr_parser_preserves_low_confidence_and_review_status() -> None:
    image = Image.new("RGB", (100, 40), "white")
    output = BytesIO()
    image.save(output, format="PNG")

    units = OCRParser(StubOCRProvider(confidence=0.42), confidence_threshold=0.8).parse(
        output.getvalue(), "scan.png"
    )

    assert units[0].extraction_method == ExtractionMethod.OCR
    assert units[0].extraction_confidence == 0.42
    assert units[0].review_status == ReviewStatus.REVIEW_REQUIRED
    assert units[0].source_locator == {"file": "scan.png", "image": 1, "row": 2}
